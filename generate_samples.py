import os
import fitz
import pandas as pd
from PIL import Image, ImageDraw
import docx

os.makedirs("demo_docs", exist_ok=True)

# 1. pump_manual.pdf (contains a long description to trigger multiple chunks)
print("Generating pump_manual.pdf with long content...")
doc = fitz.open()

# Page 1 - Large page to hold 800 words
page = doc.new_page(width=1000, height=2000)

# Generate a list of words of length 800 containing scattered asset IDs
words = []
for i in range(800):
    if i == 100:
        words.append("P-101")
    elif i == 400:
        words.append("V-200")
    elif i == 700:
        words.append("B-3")
    else:
        words.append(f"parameter_{i}")
long_description = " ".join(words)

# Insert the long description using a textbox that handles wrapping
rect = fitz.Rect(10, 10, 990, 1990)
page.insert_textbox(rect, "Centrifugal pump P-101 is designed for continuous operation. Detailed parameters:\n" + long_description)

# Page 2
page2 = doc.new_page()
page2.insert_text((50, 50), "This is page 2 content. Maintenance of V-101 and C-12 should be done every 6 months.")

doc.save("demo_docs/pump_manual.pdf")
doc.close()

# 2. maintenance_history.xlsx
print("Generating maintenance_history.xlsx...")
data_maint = {
    "Asset": ["P-101", "P-102"],
    "Date": ["2026-01-12", "2026-02-15"],
    "Issue": ["Bearing vibration", "Coupling wear"]
}
data_insp = {
    "Asset": ["P-101", "P-102"],
    "Inspector": ["Rahul", "Sarah"],
    "Status": ["OK", "Needs repair"]
}
df_maint = pd.DataFrame(data_maint)
df_insp = pd.DataFrame(data_insp)

with pd.ExcelWriter("demo_docs/maintenance_history.xlsx") as writer:
    df_maint.to_excel(writer, sheet_name="Maintenance", index=False)
    df_insp.to_excel(writer, sheet_name="Inspection", index=False)

# 3. inspection_report.jpg
print("Generating inspection_report.jpg...")
img = Image.new('RGB', (800, 400), color=(255, 255, 255))
d = ImageDraw.Draw(img)
text_lines = [
    "Drawing title: Pump Arrangement",
    "Tag: P-101",
    "Motor: M-101",
    "Flow direction: INLET -> OUTLET"
]
for i, line in enumerate(text_lines):
    d.text((20, 20 + i*40), line, fill=(0, 0, 0))
img.save("demo_docs/inspection_report.jpg")

# 4. failure_log.txt
print("Generating failure_log.txt...")
with open("demo_docs/failure_log.txt", "w", encoding="utf-8") as f:
    f.write("Log Entry: 2026-07-21 14:00:00\nSystem shutdown initiated due to overpressure on line L-102.\nOperator: John Doe\nStatus: Resolved.")

# 5. shutdown_procedure.docx
print("Generating shutdown_procedure.docx...")
doc_word = docx.Document()
doc_word.add_heading("Emergency Shutdown Procedure", 0)
doc_word.add_paragraph("1. Cut power to motor M-101.")
doc_word.add_paragraph("2. Close inlet valve V-101.")
doc_word.add_paragraph("3. Verify pressure gauge PG-101 reads zero.")
doc_word.save("demo_docs/shutdown_procedure.docx")

print("All sample files generated successfully!")
